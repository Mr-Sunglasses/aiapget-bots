import logging
import os
import re
import tempfile
from pathlib import Path

import openpyxl
from docx import Document
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import Application, CommandHandler, ContextTypes, MessageHandler, filters

load_dotenv()

LOG_FORMAT = "%(asctime)s %(levelname)s %(name)s %(message)s"
logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)
LOGGER = logging.getLogger("cleanerbot")

# Keycap digit emojis: digit + optional \uFE0F (variation selector) + \u20E3 (combining enclosing keycap)
# Matches 0️⃣ through 9️⃣ and also bare digit+keycap like 3⃣
KEYCAP_EMOJI = r"[\d#*]\uFE0F?\u20E3"

NOISE_PATTERN = re.compile(
    r"(?:"
    rf"(?:{KEYCAP_EMOJI})+"
    r"\s*"
    r"(?:\.{2,}|…)?"
    r"\s*"
    r"(?:\[\d+/\d+\])?"
    r"\s*"
    r"(?:@[A-Za-z0-9_]+)?"
    r"\s*"
    r")"
)


def clean_text(text: str) -> str:
    """Remove keycap-emoji noise prefixes from text."""
    if not text:
        return text
    cleaned, n = NOISE_PATTERN.subn("", text)
    if n > 0:
        return cleaned.strip()
    return text


def clean_xlsx(input_path: Path, output_path: Path) -> int:
    """Clean noise from all string cells in an xlsx file. Returns count of cleaned cells."""
    wb = openpyxl.load_workbook(input_path)
    cleaned_count = 0
    for ws in wb:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    original = cell.value
                    cleaned = clean_text(original)
                    if cleaned != original:
                        cell.value = cleaned
                        cleaned_count += 1
    wb.save(output_path)
    return cleaned_count


def clean_docx(input_path: Path, output_path: Path) -> int:
    """Clean noise from all paragraphs in a docx file. Returns count of cleaned paragraphs."""
    doc = Document(input_path)
    cleaned_count = 0

    for para in doc.paragraphs:
        full_text = para.text
        cleaned = clean_text(full_text)
        if cleaned == full_text:
            continue

        cleaned_count += 1
        # Rebuild runs: clear all run texts, put cleaned text in the first run
        runs = para.runs
        if not runs:
            continue
        remaining = cleaned
        for i, run in enumerate(runs):
            if i == 0:
                run.text = remaining
            else:
                run.text = ""

    doc.save(output_path)
    return cleaned_count


def clean_docx_tables(input_path: Path, output_path: Path) -> int:
    """Clean noise from table cells in a docx file that was already partially cleaned."""
    doc = Document(input_path)
    cleaned_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    cleaned = clean_text(full_text)
                    if cleaned == full_text:
                        continue
                    cleaned_count += 1
                    runs = para.runs
                    if not runs:
                        continue
                    for i, run in enumerate(runs):
                        if i == 0:
                            run.text = cleaned
                        else:
                            run.text = ""
    doc.save(output_path)
    return cleaned_count


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Send me a .docx or .xlsx file and I'll remove the noise "
        "(keycap emojis, [N/M] numbering, @mentions) from question text "
        "and send the cleaned file back."
    )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    doc = update.message.document
    if not doc:
        return

    file_name = doc.file_name or "unknown"
    suffix = Path(file_name).suffix.lower()

    if suffix not in (".docx", ".xlsx"):
        await update.message.reply_text(
            f"Unsupported file type: {suffix}\nPlease send a .docx or .xlsx file."
        )
        return

    await update.message.reply_text(f"Processing {file_name} ...")

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_dir_path = Path(tmp_dir)
        input_path = tmp_dir_path / file_name
        output_path = tmp_dir_path / f"cleaned_{file_name}"

        tg_file = await doc.get_file()
        await tg_file.download_to_drive(input_path)

        try:
            if suffix == ".xlsx":
                cleaned_count = clean_xlsx(input_path, output_path)
            else:
                cleaned_count = clean_docx(input_path, output_path)
                cleaned_count += clean_docx_tables(output_path, output_path)
        except Exception:
            LOGGER.exception("Error cleaning %s", file_name)
            await update.message.reply_text(f"Error processing {file_name}. Is it a valid {suffix} file?")
            return

        with open(output_path, "rb") as f:
            await update.message.reply_document(
                document=f,
                filename=file_name,
                caption=f"Cleaned {cleaned_count} cell(s)/paragraph(s).",
            )


def main() -> None:
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    if not token:
        raise RuntimeError("TELEGRAM_BOT_TOKEN is not set.")

    application = Application.builder().token(token).build()
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    LOGGER.info("Cleaner Bot started.")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
