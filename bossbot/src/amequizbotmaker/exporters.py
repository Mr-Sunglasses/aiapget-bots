"""Generate XLSX and DOCX files from quiz JSON data.

Ported from quiz-exporter with adjustments for the unified bot.
"""

from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document
from docx.shared import Pt


OPTION_LETTERS = "ABCDEFGHIJ"


def _find_template_dir() -> Path:
    """Locate the ``templates/`` directory.  Checks several locations."""
    candidates = [
        # Running from the repo root (normal dev / Docker with COPY .)
        Path(__file__).resolve().parent.parent.parent.parent / "templates",
        # Docker: /app/templates
        Path("/app/templates"),
        # CWD fallback
        Path("templates"),
    ]
    for candidate in candidates:
        if candidate.is_dir():
            return candidate
    return candidates[0]  # fallback


TEMPLATE_DIR = _find_template_dir()
XLSX_TEMPLATE = TEMPLATE_DIR / "questions_output.xlsx"
DOCX_TEMPLATE = TEMPLATE_DIR / "question_output.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_quiz_json(data: list | dict) -> tuple[str, str, list[dict]]:
    """Return ``(file_name, title, polls)`` from the quiz JSON.

    Accepts both the legacy array format and the newer dict format::

        # Legacy: [title_str, {poll}, {poll}, ...]
        # New:    {"file_name": "...", "data": [title_str, {poll}, ...]}
    """
    file_name = ""
    inner = data

    if isinstance(data, dict):
        file_name = data.get("file_name", "")
        inner = data.get("data", [])

    title = ""
    polls: list[dict] = []
    for item in inner:
        if isinstance(item, str):
            title = item
        elif isinstance(item, dict):
            polls.append(item)
    return file_name, title, polls


def get_file_name(data: list | dict) -> str:
    """Extract ``file_name`` from the quiz JSON, or return ``""``."""
    file_name, _title, _polls = _parse_quiz_json(data)
    return file_name


def get_poll_count(data: list | dict) -> int:
    """Count the number of poll objects in the quiz JSON."""
    _file_name, _title, polls = _parse_quiz_json(data)
    return len(polls)


def _right_answer_letter(correct_option_id: int | None) -> str:
    if correct_option_id is None:
        return ""
    if 0 <= correct_option_id < len(OPTION_LETTERS):
        return OPTION_LETTERS[correct_option_id]
    return str(correct_option_id + 1)


def _right_answer_number(correct_option_id: int | None) -> int | str:
    """Return 1-based numeric answer for XLSX (A=1, B=2, …)."""
    if correct_option_id is None:
        return ""
    return correct_option_id + 1


def _question_type(poll: dict) -> str:
    if poll.get("allows_multiple_answers"):
        return "MULTICORRECT"
    return "SINGLECORRECT"


# ---------------------------------------------------------------------------
# XLSX export
# ---------------------------------------------------------------------------

_HEADER_FILL = PatternFill(start_color="004472C4", end_color="004472C4", fill_type="solid")
_HEADER_FONT = Font(bold=True, color="00FFFFFF")
_HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
_DATA_ALIGN = Alignment(vertical="top", wrap_text=True)

_HEADERS = [
    "S No.", "SUBJECT", "TOPIC", "TAGS", "QUESTION TYPE",
    "QUESTION TEXT",
    "OPTION1", "OPTION2", "OPTION3", "OPTION4",
    "OPTION5", "OPTION6", "OPTION7", "OPTION8",
    "OPTION9", "OPTION10",
    "RIGHT ANSWER", "EXPLANATION",
    "CORRECT MARKS", "NEGATIVE MARKS", "DIFFICULTY",
]


def generate_xlsx(data: list | dict, output_path: Path) -> Path:
    """Write quiz data into an XLSX file following the template format."""
    _file_name, title, polls = _parse_quiz_json(data)

    if XLSX_TEMPLATE.exists():
        wb = openpyxl.load_workbook(XLSX_TEMPLATE)
        ws = wb.active
        if ws.max_row > 1:
            ws.delete_rows(2, ws.max_row - 1)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Questions"
        for col_idx, header in enumerate(_HEADERS, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = _HEADER_FILL
            cell.font = _HEADER_FONT
            cell.alignment = _HEADER_ALIGN

    for row_idx, poll in enumerate(polls, start=2):
        options = poll.get("options") or []
        option_texts = [opt.get("text", "") for opt in options]
        while len(option_texts) < 10:
            option_texts.append(None)

        row_data = [
            row_idx - 1,                                              # S No.
            None,                                                     # SUBJECT
            None,                                                     # TOPIC
            title or None,                                            # TAGS
            _question_type(poll),                                     # QUESTION TYPE
            poll.get("question", ""),                                 # QUESTION TEXT
            *option_texts[:10],                                       # OPTION1..10
            _right_answer_number(poll.get("correct_option_id")),      # RIGHT ANSWER
            poll.get("explanation", ""),                               # EXPLANATION
            4,                                                        # CORRECT MARKS
            1,                                                        # NEGATIVE MARKS
            "Medium",                                                 # DIFFICULTY
        ]

        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = _DATA_ALIGN

    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

_MAX_DOCX_OPTIONS = 5


def generate_docx(data: list | dict, output_path: Path) -> Path:
    """Write quiz data into a DOCX file matching the template format."""
    _file_name, title, polls = _parse_quiz_json(data)

    doc = Document()

    # Section header
    section_para = doc.add_paragraph()
    run_prefix = section_para.add_run("###Section ")
    run_prefix.font.name = "Courier New"
    run_prefix.font.size = Pt(12)

    run_section = section_para.add_run("AYURVEDA")
    run_section.bold = True
    run_section.font.name = "Times"
    run_section.font.size = Pt(12)

    doc.add_paragraph("")

    for idx, poll in enumerate(polls, start=1):
        doc.add_paragraph(" #English_directions")
        doc.add_paragraph(f" #Question {idx} ")
        doc.add_paragraph("")

        doc.add_paragraph(poll.get("question", ""))

        options = poll.get("options") or []
        for opt_idx in range(_MAX_DOCX_OPTIONS):
            letter = OPTION_LETTERS[opt_idx]
            opt_text = ""
            if opt_idx < len(options):
                opt_text = options[opt_idx].get("text", "")

            if opt_idx == 0:
                doc.add_paragraph(f"#Options ###{letter}")
            else:
                doc.add_paragraph(f"###{letter}")
            doc.add_paragraph(opt_text)

        doc.add_paragraph("#Correct_option")
        doc.add_paragraph(_right_answer_letter(poll.get("correct_option_id")))

        doc.add_paragraph("#Solution")
        explanation = poll.get("explanation", "")
        doc.add_paragraph(
            explanation if explanation else "As per reference in the original question."
        )

        doc.add_paragraph("#tag")
        doc.add_paragraph(title if title else "Topic name")

        doc.add_paragraph("")

    doc.save(output_path)
    return output_path
